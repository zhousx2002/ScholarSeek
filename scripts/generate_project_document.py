# -*- coding: utf-8 -*-
from __future__ import annotations

import html
from pathlib import Path

from PIL import Image as PILImage, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "submission_docs"
ASSET_DIR = OUTPUT_DIR / "assets"
PROJECT_NAME = "ScholarSeek：面向复杂科研查询的多阶段智能论文搜索与推荐系统"
SHORT_NAME = "ScholarSeek"
TEAM_NAME = "[待填写：团队名称]"
GROUP_NAME = "华为赛道·赛题三"
VERSION = "V1.0"
DATE = "2026.06.23"
DOCX_PATH = OUTPUT_DIR / "待填写团队_ScholarSeek_项目文档.docx"
PDF_PATH = OUTPUT_DIR / "待填写团队_ScholarSeek_项目文档.pdf"
CHINESE_FONT = Path(r"C:\Windows\Fonts\msyh.ttc")


def build_elements():
    architecture = ASSET_DIR / "architecture.png"
    spar_flow = ASSET_DIR / "spar_flow.png"
    training_flow = ASSET_DIR / "training_flow.png"
    return [
        ("h1", "1 项目概况"),
        ("h2", "1.1 背景和基础"),
        ("p", "科研文献检索是选题论证、技术调研和实验设计的基础环节。华为赛道赛题三“科研场景下复杂学术查询的智能论文搜索与推荐”指出，真实科研查询通常同时包含研究主题、方法、数据集、时间范围、发表场所和结果类型等多维约束，传统关键词搜索难以完成语义理解、多轮检索、引文网络探索与跨文献关联推理。赛题要求系统端到端完成查询理解、策略迭代、候选过滤、论文综合排序和结构化归纳，并以F1、运行效率和结果结构化程度作为核心评价维度。"),
        ("p", "ScholarSeek由此立项，目标是构建一套可运行、可训练、可评测的学术搜索Agent。项目借鉴PaSa的Crawler/Selector思想与SPAR的RefChain、Query Evolution机制，但未直接复制其工程，而是在轻量Python服务中重新实现多源检索、单层引文扩展、三态相关性判断、可训练重排序和证据约束回答。当前系统已经具备Web交互界面、标准/增强双模式、OpenAlex、Semantic Scholar、arXiv三类检索后端、Qwen自然语言规划与回答、PaSa训练数据管线及端到端评测工具。"),
        ("p", "已有工作基础包括：完成25项自动化测试；构建33,551条AutoScholarQuery训练记录与609,625个query-paper样本对；训练并集成体积为724字节的compact fallback reranker，其开发集AUC为0.8035；完成BAAI/bge-reranker-base主模型的hard-negative与RankNet pairwise训练代码、GPU混合精度烟雾验证和在线加载验证。主模型正式训练与公开/隐藏测试集最终指标将在提交前固化，本文对未完成指标统一标注为“待测”。"),
        ("table", ["角色", "人员", "主要职责"], [
            ["项目负责人", "[待填写]", "总体方案、系统架构、进度与答辩"],
            ["算法成员", "[待填写]", "查询规划、重排序训练、评测"],
            ["工程成员", "[待填写]", "检索API、后端服务、前端与部署"],
        ], [3.0, 3.0, 10.0]),
        ("h2", "1.2 场景和价值"),
        ("p", "系统面向高校研究生、科研人员、企业研发工程师及科技情报人员，适用于开题调研、综述撰写、技术路线比较、前沿追踪、竞品分析和跨领域知识发现。用户无需掌握数据库检索语法，只需输入自然语言问题，系统即可拆解意图、自动构造检索式、跨源召回、探索引文关系并输出带相关性分数和证据说明的论文列表。"),
        ("table", ["典型场景", "用户痛点", "ScholarSeek价值"], [
            ["开题与综述", "查询宽泛、关键词不稳定、人工筛选耗时", "分解主题/方法/约束，多轮扩展并聚合论文"],
            ["技术路线比较", "需同时覆盖方法、应用、局限与替代方案", "Query Evolution从高相关论文生成互补方向"],
            ["前沿追踪", "需兼顾新颖性、权威性和相关性", "年份、引用、来源与语义相关性综合排序"],
            ["跨领域检索", "术语差异导致漏检", "Qwen改写与OpenAlex跨学科覆盖"],
            ["证据问答", "通用LLM容易脱离论文证据", "仅基于已检索论文生成带编号引用的回答"],
        ], [3.0, 5.0, 8.0]),
        ("p", "相较传统搜索引擎，ScholarSeek的核心增量不在于替代数据库，而在于提供“自然语言意图到可验证论文集合”的自动化工作流。相较PaSa，项目增加了可控的一层RefChain、前端过程可视化和轻量fallback；相较SPAR，项目将树式深层扩展约束为单层有界并行，减少API调用与级联噪声；相较仅依赖LLM的Agent，系统以确定性规则、训练模型和缓存共同控制成本与稳定性。"),
        ("table", ["方案", "查询理解", "引文扩展", "可训练精排", "成本控制", "本项目采用情况"], [
            ["传统关键词检索", "弱", "无", "平台内置", "高", "作为API基础能力"],
            ["PaSa", "LLM Agent", "Crawler迭代", "Selector/RL", "中", "借鉴Crawler/Selector分工"],
            ["SPAR", "查询分解/演化", "RefChain", "多阶段判断", "中", "迁移Query Evolution与单层RefChain"],
            ["ScholarSeek", "Qwen+规则兜底", "单层references/citations", "BGE+hard negative", "Fast/SPAR双模式", "已实现"],
        ], [2.7, 2.8, 3.0, 3.0, 2.8, 3.7]),
        ("h2", "1.3 所需支持"),
        ("p", "项目开发与训练使用Windows工作站、NVIDIA GeForce GTX 1660 Ti 6GB显存和E盘Miniconda环境。在线推理需要网络访问OpenAlex、Semantic Scholar、arXiv与阿里云百炼DashScope。Semantic Scholar API Key可提升引文接口稳定性；OpenAlex配置联系邮箱进入polite pool；Qwen Key仅保存在.gitignore排除的.env文件中。"),
        ("bullet", "算力：主重排序模型训练建议使用16GB以上显存GPU；现有6GB显卡通过batch size 1、FP16、梯度累积和gradient checkpointing完成训练。"),
        ("bullet", "存储：原始PaSa/SPAR数据约17MB，hard-negative训练对约162MB，BGE基础权重与微调权重各约1.1GB。"),
        ("bullet", "接口：OpenAlex与arXiv可匿名使用；Semantic Scholar建议申请Key；Qwen使用qwen-plus兼容接口。"),
        ("bullet", "提交：其他材料压缩包采用白名单构建，排除.env、数据集、Conda环境和完整权重，自动检查不超过200MB。"),
        ("pagebreak",),
        ("h1", "2 项目规划"),
        ("h2", "2.1 整体目标"),
        ("p", "项目总体目标是交付一个可现场演示、可通过命令行复现、可在PaSa与SPAR benchmark上自动评测的端到端论文搜索系统。系统既要覆盖复杂查询的召回，又要避免无界Agent循环造成的延迟和费用，在竞赛F1、效率与结构化展示三项指标间取得平衡。"),
        ("table", ["目标维度", "目标内容", "验收方式"], [
            ["功能完整性", "自然语言查询、改写、跨源召回、RefChain、精排、回答、导出", "Web现场演示与API响应"],
            ["检索质量", "提高Top-K精确率与召回率，控制无关论文", "Precision/Recall/F1/Hit@K/MRR"],
            ["运行效率", "Fast模式快速首屏；SPAR模式限制深度、种子和并发", "API次数、Token和端到端延时"],
            ["泛化能力", "覆盖对话式、长约束、中文与英文查询", "AutoScholarQuery、RealScholarQuery、SPAR bench"],
            ["结构化展示", "论文列表、分数、来源、年份、摘要、选择与导出", "前端可用性检查"],
            ["可交付性", "源代码包低于200MB且不泄露密钥", "自动打包脚本与文件审计"],
        ], [3.0, 8.0, 5.0]),
        ("h2", "2.2 技术创新点"),
        ("p", "（1）成本感知的Fast/SPAR双模式。Fast模式采用启发式分解、OpenAlex快速召回与缓存，优先响应速度；SPAR模式启用Qwen规划、Judgement Agent、引文扩展与Query Evolution，优先复杂查询的覆盖率。用户可在同一界面切换，系统在响应中返回pipeline_trace，量化每阶段候选贡献。"),
        ("p", "（2）受控单层RefChain。SPAR的引文探索能够提升召回，但深层链式扩展容易放大噪声与成本。本项目只对高相关论文进行一层references/citations扩展，默认最多3个种子、每种子最多8篇，并采用有界线程池并行。无Semantic Scholar Key时仅尝试references，429或超时自动降级，不阻断主链路。"),
        ("p", "（3）面向困难负样本的语义重排序。训练集不再仅随机抽取无关标题，而是通过倒排索引和IDF从33,267个标题中挖掘词面相似但标签错误的hard negatives。例如查询包含image retrieval时，将同样包含image/retrieval但研究目标不同的论文作为负例。主模型采用BAAI/bge-reranker-base，并以RankNet损失直接学习正论文得分高于困难负论文。"),
        ("formula", "L_rank = -log σ(s(q, p⁺) - s(q, p⁻))"),
        ("p", "（4）混合Judgement Agent。系统先使用cross-encoder或compact模型计算相关性，将候选划分为related、uncertain、unrelated；仅在SPAR-Qwen模式对边界候选进行LLM复核，从而避免逐篇调用大模型。只有related论文能够成为引文扩展与查询演化种子，防止搜索漂移。"),
        ("p", "（5）证据约束生成与可解释轨迹。回答模块只接收Top-K论文的标题、摘要、年份、引用与URL，要求Qwen按[1][2]引用；若证据不足则明确给出search gaps。前端同步展示reranker score、判断理由与pipeline trace，使结果可检查、可复现。"),
        ("image", str(architecture), "图2-1 ScholarSeek系统分层架构"),
        ("h3", "架构设计原则"),
        ("bullet", "模块解耦：Planner、Retriever、Judgement、Reranker和Answer Synthesizer通过统一数据结构协作，可单独替换模型或数据源。"),
        ("bullet", "质量优先：复杂查询走SPAR-Qwen与语义cross-encoder，相关性判断始终以原始用户意图为基准。"),
        ("bullet", "成本有界：RefChain深度、种子数量、演化query数、并发数和回答证据数均设置硬上限。"),
        ("bullet", "可观测与可降级：pipeline trace记录阶段贡献；任一外部API或模型不可用时均有明确回退路径。"),
        ("pagebreak",),
        ("h1", "3 实施方案"),
        ("h2", "3.1 技术可行性分析"),
        ("p", "数据可行性方面，项目使用PaSa AutoScholarQuery训练/验证/测试集、RealScholarQuery真实查询集，以及SPAR的AutoScholarQuery_test和spar_bench。训练集与测试集严格分离；SPAR benchmark只用于端到端测试，不参与参数训练。当前统一数据目录为E:\\DATASET\\PasaDataSet。"),
        ("table", ["数据集", "用途", "记录数", "是否参与训练"], [
            ["AutoScholarQuery/train", "重排序训练", "33,551", "是"],
            ["AutoScholarQuery/dev", "模型选择", "1,000", "仅验证"],
            ["AutoScholarQuery/test", "自动测试", "1,000", "否"],
            ["RealScholarQuery/test", "真实查询泛化", "50", "否"],
            ["benchmark/AutoScholarQuery_test", "SPAR对比", "1,000", "否"],
            ["benchmark/spar_bench", "复杂人工查询", "50", "否"],
        ], [5.5, 4.5, 2.5, 3.5]),
        ("p", "知识与检索可行性方面，OpenAlex覆盖跨学科论文元数据，Semantic Scholar提供计算机领域检索与引文图接口，arXiv补充最新预印本。三者均被封装为统一Paper结构，按照DOI与规范化标题去重，并合并摘要、引用数、作者、venue与开放链接。单一API故障时系统记录warning并继续其他来源。"),
        ("p", "算力可行性方面，compact reranker可在CPU上即时运行；BGE主模型在6GB显存上已经完成FP16、gradient checkpointing、pairwise forward、保存与在线加载烟雾验证。正式训练通过max-train-pairs进行分阶段扩展，并按开发集loss保存最佳权重。工程可行性方面，后端仅使用Python标准HTTP服务和少量模型依赖，前端为原生HTML/CSS/JavaScript，部署门槛低。"),
        ("h2", "3.2 技术细节"),
        ("h3", "3.2.1 总体架构与数据流"),
        ("p", "系统划分为交互层、Agent编排层、检索层、判断与排序层、生成层和评测层。Web端向/api/search发送query、strategy、sources和Top-K参数；SearchService生成QueryPlan，MultiSourceRetriever并发/顺序访问学术API并统一元数据；Fast模式直接精排，SPAR模式进入RefChainPlanner；最终结果缓存10分钟并异步调用/api/answer生成回答。"),
        ("image", str(spar_flow), "图3-1 SPAR增强检索流程"),
        ("h3", "3.2.2 查询理解与演化"),
        ("p", "启发式Query Planner负责年份范围、技术短语、核心词和可选词提取，并清除what works、can you tell me等对话模板词。Qwen Planner使用JSON Schema输出search_queries、must_terms、optional_terms、year_from与year_to；当Qwen不可用时自动回退启发式方案。Query Evolver读取原始query、已搜索query和最多5篇高相关论文，分别生成方法比较、应用实现和局限挑战三个方向的新检索式，并做规范化去重。"),
        ("h3", "3.2.3 多源召回与元数据融合"),
        ("p", "OpenAlex客户端对自然语言标点进行清洗，按年份生成filter；Semantic Scholar请求paperId、title、abstract、year、venue、authors、externalIds、citationCount等字段；arXiv解析Atom Feed并执行年份过滤。合并阶段优先使用DOI键，否则使用小写字母数字化标题键；重复论文保留更高引用数并补齐摘要、venue、作者、DOI和URL。"),
        ("h3", "3.2.4 RefChain与Judgement Agent"),
        ("p", "RefChainPlanner先对初检候选排序与判断，选择related论文作为种子；CitationExpander依据Semantic Scholar paperId、DOI或arXiv ID调用/graph/v1/paper/{paper_id}/references和/citations。引文候选再次通过Judgement Agent过滤；随后Query Evolver生成新query并有界并行检索。所有候选最终去重并针对原始query重新排序。该流程max_refchain_depth固定为1。"),
        ("p", "Judgement Agent默认阈值为related≥0.62、unrelated<0.32，中间样本标记uncertain。评分优先使用主cross-encoder，加载失败时使用compact fallback，再失败时使用词项覆盖规则。SPAR-Qwen模式可对边界样本返回decision、score与reason。"),
        ("h3", "3.2.5 重排序模型与训练"),
        ("image", str(training_flow), "图3-2 PaSa hard-negative pairwise训练流程"),
        ("p", "训练数据生成器从AutoScholarQuery答案构造正样本，每个正样本配置2个随机负样本和4个hard negatives。HardNegativeMiner为标题建立倒排索引，使用IDF与Jaccard相关信号寻找高相似错误论文，避免O(N×M)全池扫描。正式数据包括609,625个query-paper pair，其中可形成约522,520个正负triplet。"),
        ("p", "主模型BAAI/bge-reranker-base以query与title为双句输入，输出标量logit。训练默认使用pairwise RankNet loss、FP16、梯度累积16、gradient checkpointing、max_length 256和AdamW；开发集计算pairwise accuracy与loss，并只保存dev loss最优模型。在线推理默认batch size 4，读取训练元数据中的max_length，在GPU上使用autocast。"),
        ("h3", "3.2.6 排序、回答与前端"),
        ("p", "无训练模型时，规则排序综合核心词覆盖、扩展词覆盖、标题命中、短语命中、年份适配和引用奖励；有主模型时，cross-encoder覆盖规则分数。Qwen Answer Synthesizer只使用Top-8论文证据，输出direct answer、highly relevant papers、partial/needs verification和search gaps。前端支持Fast/SPAR切换、年份筛选、复选、删除、详情展开、JSON下载和结果复制。"),
        ("h3", "3.2.7 评测方案"),
        ("p", "端到端评测脚本直接读取PaSa或SPAR JSONL，对预测Top-K标题与gold标题进行规范化精确匹配。单query指标定义如下，其中P为预测集合，G为标注集合："),
        ("formula", "Precision = |P∩G| / |P|；Recall = |P∩G| / |G|；F1 = 2PR / (P+R)"),
        ("p", "同时报告macro Precision/Recall/F1、micro Precision/Recall/F1、Hit@K、MRR、平均端到端延时与pipeline trace。评测分别覆盖AutoScholarQuery test、RealScholarQuery、SPAR AutoScholarQuery_test和spar_bench。竞赛最终结果在主模型训练完成后固化，禁止以开发集结果替代测试集成绩。"),
        ("table", ["指标/状态", "当前值", "说明"], [
            ["自动化测试", "25项通过", "代码回归测试，非检索F1"],
            ["Compact dev AUC", "0.8035", "PaSa pair开发集，fallback模型"],
            ["Hard-negative准备", "609,625 pairs", "训练33,551 queries"],
            ["SPAR烟雾延时", "35.22s→12.41s", "单源小规模参数，有界并行优化；非竞赛最终值"],
            ["主模型test F1", "待测", "BGE正式训练完成后执行"],
            ["RealScholarQuery F1", "待测", "严格保留测试集"],
        ], [4.0, 3.5, 8.5]),
        ("h3", "3.2.8 效率、可靠性与安全"),
        ("bullet", "缓存：检索结果按query、策略、来源、参数和reranker路径缓存10分钟。"),
        ("bullet", "并行：引文请求最多4线程，演化query最多3线程；所有循环有深度和数量上限。"),
        ("bullet", "降级：Qwen失败回退启发式；Semantic Scholar失败继续OpenAlex/arXiv；主模型失败回退compact。"),
        ("bullet", "安全：API Key仅在.env中读取，/api/config只返回掩码，提交脚本明确排除.env。"),
        ("bullet", "可交付：scripts/build_submission.ps1采用白名单打包并强制校验200MB上限。"),
        ("h2", "3.3 计划和分工"),
        ("table", ["阶段", "工作内容", "状态/交付物"], [
            ["阶段一：需求与基线", "解析赛题、PaSa/SPAR调研、多源API接入", "已完成"],
            ["阶段二：原型系统", "Web前端、Fast检索、Qwen回答、交互功能", "已完成"],
            ["阶段三：增强检索", "Query Evolution、RefChain、Judgement、pipeline trace", "已完成"],
            ["阶段四：模型训练", "hard negative、BGE pairwise训练、模型选择", "进行中"],
            ["阶段五：公开评测", "四类测试集、消融、效率与错误分析", "待完成"],
            ["阶段六：提交与答辩", "文档、视频、其他.zip、演示脚本", "进行中"],
        ], [4.0, 7.5, 4.5]),
        ("table", ["成员", "分工", "成果"], [
            ["[待填写：负责人]", "总体设计、Agent编排、文档答辩", "系统架构与项目管理"],
            ["[待填写：算法]", "数据构造、重排序训练、评测", "模型与实验报告"],
            ["[待填写：工程]", "API、前后端、部署与视频", "可演示系统与提交包"],
        ], [4.0, 6.0, 6.0]),
        ("pagebreak",),
        ("h1", "4 参考资料"),
        ("ref", "[1] He Y, Huang G, Feng P, et al. PaSa: An LLM Agent for Comprehensive Academic Paper Search. ACL 2025, arXiv:2501.10120."),
        ("ref", "[2] Ajith A, Xia M, Chevalier A, et al. LitSearch: A Retrieval Benchmark for Scientific Literature Search. EMNLP 2024, arXiv:2407.18940."),
        ("ref", "[3] Feldman S, et al. AstaBench: Rigorous Benchmarking of AI Agents with a Scientific Research Suite. arXiv:2510.21652, 2025."),
        ("ref", "[4] Shi X, Li Y, Kou Q, et al. SPAR: Scholar Paper Retrieval with LLM-based Agents for Enhanced Academic Search. arXiv:2507.15245, 2025."),
        ("ref", "[5] Skarlinski M, et al. Language Agents for Answering Questions from Scientific Literature. NeurIPS 2024."),
        ("ref", "[6] Khattab O, Santhanam K, Li X, et al. Demonstrate-Search-Predict: Composing Retrieval and Language Models for Knowledge-Intensive NLP. arXiv:2212.14024, 2022."),
        ("ref", "[7] Feng P, He Y, Huang G, et al. AGILE: A Novel Framework of LLM Agents. NeurIPS 2024."),
        ("ref", "[8] Muennighoff N, et al. GritLM: Generative Representational Instruction Tuning. arXiv:2402.09906, 2024."),
        ("ref", "[9] Press O, Zhang M, Min S, et al. Measuring and Narrowing the Compositionality Gap in Language Models. arXiv:2210.03350, 2022."),
        ("ref", "[10] Lee D, Sohn S S, Lee B, et al. Domain-aligned LLM Framework for Trustworthy Scientific Q/A via Query Reformulation RAG. ChemRxiv, 2025."),
        ("ref", "[11] OpenAlex API Documentation. https://docs.openalex.org/"),
        ("ref", "[12] Semantic Scholar Academic Graph API. https://api.semanticscholar.org/api-docs/graph"),
        ("ref", "[13] arXiv API User's Manual. https://info.arxiv.org/help/api/"),
        ("ref", "[14] BAAI. bge-reranker-base model card. https://huggingface.co/BAAI/bge-reranker-base"),
    ]


def font(size, bold=False):
    return ImageFont.truetype(str(CHINESE_FONT), size=size, index=0)


def rounded_box(draw, xy, text, fill, outline, width=3, text_fill="#102456", title_size=27, body=None):
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=width)
    x1, y1, x2, y2 = xy
    bbox = draw.textbbox((0, 0), text, font=font(title_size, True))
    draw.text(((x1 + x2 - bbox[2]) / 2, y1 + 18), text, font=font(title_size, True), fill=text_fill)
    if body:
        for idx, line in enumerate(body):
            draw.text((x1 + 18, y1 + 62 + idx * 31), "• " + line, font=font(19), fill="#303846")


def arrow(draw, start, end, color="#557dff", width=5):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 14, y - 9), (x - 14, y + 9)], fill=color)


def create_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    canvas = PILImage.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(canvas)
    d.text((70, 35), "ScholarSeek 分层架构", font=font(42, True), fill="#081847")
    layers = [
        ("交互层", ["Fast / SPAR模式", "查询、筛选、导出"], "#eef5ff"),
        ("Agent编排层", ["Qwen Planner", "RefChain / Evolution"], "#edf9f2"),
        ("多源检索层", ["OpenAlex", "Semantic Scholar / arXiv"], "#fff6e6"),
        ("判断与排序层", ["Judgement Agent", "BGE / compact fallback"], "#f5efff"),
        ("生成与展示层", ["证据约束Qwen回答", "列表、摘要、轨迹"], "#fff0f2"),
    ]
    y = 130
    for i, (title, body, color) in enumerate(layers):
        rounded_box(d, (160, y, 1440, y + 120), title, color, "#bfd0f2", body=body)
        if i < len(layers) - 1:
            d.line([(800, y + 120), (800, y + 151)], fill="#557dff", width=5)
            d.polygon([(800, y + 158), (790, y + 143), (810, y + 143)], fill="#557dff")
        y += 145
    canvas.save(ASSET_DIR / "architecture.png")

    canvas = PILImage.new("RGB", (1800, 760), "white")
    d = ImageDraw.Draw(canvas)
    d.text((60, 30), "SPAR 增强检索（有界单层）", font=font(40, True), fill="#081847")
    boxes = [
        ("自然语言Query", ["主题/方法/约束"]),
        ("初始多源检索", ["去重与元数据融合"]),
        ("Judgement", ["related/uncertain/unrelated"]),
        ("双路扩展", ["RefChain一层", "Query Evolution"]),
        ("再次判断", ["过滤扩展噪声"]),
        ("最终精排", ["Top-K + Qwen回答"]),
    ]
    for i, (title, body) in enumerate(boxes):
        x = 45 + i * 292
        rounded_box(d, (x, 250, x + 245, 440), title, "#f6f9ff", "#9eb8e8", title_size=23, body=body)
        if i < len(boxes) - 1:
            arrow(d, (x + 245, 345), (x + 284, 345))
    d.text((505, 520), "仅related论文成为种子；深度=1；种子≤3；失败自动降级", font=font(24, True), fill="#325caa")
    canvas.save(ASSET_DIR / "spar_flow.png")

    canvas = PILImage.new("RGB", (1700, 780), "white")
    d = ImageDraw.Draw(canvas)
    d.text((60, 35), "PaSa 语义重排序训练流程", font=font(40, True), fill="#081847")
    boxes = [
        ("PaSa Query", ["33,551条训练查询"]),
        ("正样本", ["标注答案论文"]),
        ("负样本构造", ["2个随机负例", "4个hard negatives"]),
        ("BGE双句编码", ["query + positive", "query + negative"]),
        ("RankNet Loss", ["拉大正负分差"]),
        ("最佳权重", ["按dev loss保存"]),
    ]
    for i, (title, body) in enumerate(boxes):
        x = 30 + i * 278
        rounded_box(d, (x, 230, x + 230, 440), title, "#f8fbff", "#aac2ee", title_size=22, body=body)
        if i < len(boxes) - 1:
            arrow(d, (x + 230, 335), (x + 270, 335))
    d.text((430, 535), "FP16 + 梯度累积 + Gradient Checkpointing", font=font(27, True), fill="#325caa")
    d.text((610, 600), "L = -log σ(s⁺ - s⁻)", font=font(30, True), fill="#081847")
    canvas.save(ASSET_DIR / "training_flow.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **borders):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        settings = borders.get(edge)
        if settings is None:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in settings.items():
            element.set(qn("w:{}".format(key)), str(value))


def apply_three_line_table_style(table):
    no_border = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    top_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid_border = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    bottom_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=no_border,
                left=no_border,
                bottom=no_border,
                right=no_border,
                insideH=no_border,
                insideV=no_border,
            )

    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top_border, bottom=mid_border, left=no_border, right=no_border)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=bottom_border, left=no_border, right=no_border)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_docx(document):
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [("Title", 25, "081847"), ("Heading 1", 18, "081847"), ("Heading 2", 14, "102456"), ("Heading 3", 12, "325CAA")]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    header = section.header.paragraphs[0]
    header.text = SHORT_NAME + " 项目文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(124, 131, 145)
    add_page_number(section.footer.paragraphs[0])


def add_docx_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Cm(width)
    apply_three_line_table_style(table)
    document.add_paragraph()


def create_docx(elements):
    document = Document()
    configure_docx(document)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    run = p.add_run("第八届中国研究生人工智能创新大赛")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Microsoft YaHei"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    run = p.add_run(PROJECT_NAME)
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(8, 24, 71)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目文档").bold = True
    p.runs[0].font.size = Pt(22)
    for text in (VERSION, DATE, TEAM_NAME, GROUP_NAME):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(13)
        if "待填写" in text:
            run.font.color.rgb = RGBColor(196, 60, 60)
    document.add_page_break()
    document.add_heading("目录", level=1)
    for line in ["1 项目概况", "  1.1 背景和基础", "  1.2 场景和价值", "  1.3 所需支持", "2 项目规划", "  2.1 整体目标", "  2.2 技术创新点", "3 实施方案", "  3.1 技术可行性分析", "  3.2 技术细节", "  3.3 计划和分工", "4 参考资料"]:
        p = document.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.8 if line.startswith("  ") else 0)
    document.add_page_break()
    document.add_heading("记录更改历史", level=1)
    add_docx_table(document, ["序号", "更改原因", "版本", "作者", "更改日期", "备注"], [["1", "初始编制", VERSION, TEAM_NAME, DATE, "按大赛模板生成"]], [1.2, 4.2, 1.5, 3.2, 2.5, 3.2])
    document.add_page_break()
    for item in elements:
        kind = item[0]
        if kind == "h1":
            document.add_heading(item[1], level=1)
        elif kind == "h2":
            document.add_heading(item[1], level=2)
        elif kind == "h3":
            document.add_heading(item[1], level=3)
        elif kind in {"p", "ref"}:
            p = document.add_paragraph(item[1])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.74 if kind == "p" else 0)
        elif kind == "bullet":
            p = document.add_paragraph(item[1], style="List Bullet")
        elif kind == "formula":
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item[1])
            run.bold = True
            run.font.size = Pt(11)
        elif kind == "table":
            add_docx_table(document, item[1], item[2], item[3])
        elif kind == "image":
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(item[1], width=Cm(16.0))
            cap = document.add_paragraph(item[2])
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
        elif kind == "pagebreak":
            document.add_page_break()
    document.core_properties.title = PROJECT_NAME
    document.core_properties.subject = "华为赛道赛题三项目文档"
    document.core_properties.author = TEAM_NAME
    document.save(DOCX_PATH)


def pdf_styles():
    pdfmetrics.registerFont(TTFont("YaHei", str(CHINESE_FONT)))
    base = getSampleStyleSheet()
    return {
        "normal": ParagraphStyle("CN", parent=base["BodyText"], fontName="YaHei", fontSize=9.5, leading=15, alignment=TA_JUSTIFY, spaceAfter=6, firstLineIndent=20),
        "bullet": ParagraphStyle("BulletCN", parent=base["BodyText"], fontName="YaHei", fontSize=9.2, leading=14, leftIndent=16, firstLineIndent=-8, spaceAfter=4),
        "h1": ParagraphStyle("H1CN", parent=base["Heading1"], fontName="YaHei", fontSize=18, leading=24, textColor=colors.HexColor("#081847"), spaceBefore=10, spaceAfter=10),
        "h2": ParagraphStyle("H2CN", parent=base["Heading2"], fontName="YaHei", fontSize=14, leading=20, textColor=colors.HexColor("#102456"), spaceBefore=8, spaceAfter=7),
        "h3": ParagraphStyle("H3CN", parent=base["Heading3"], fontName="YaHei", fontSize=11.5, leading=17, textColor=colors.HexColor("#325CAA"), spaceBefore=7, spaceAfter=5),
        "center": ParagraphStyle("CenterCN", parent=base["BodyText"], fontName="YaHei", fontSize=11, leading=17, alignment=TA_CENTER),
        "formula": ParagraphStyle("FormulaCN", parent=base["BodyText"], fontName="YaHei", fontSize=10.5, leading=18, alignment=TA_CENTER, textColor=colors.HexColor("#081847"), spaceBefore=6, spaceAfter=8),
        "ref": ParagraphStyle("RefCN", parent=base["BodyText"], fontName="YaHei", fontSize=8.7, leading=13, alignment=TA_LEFT, spaceAfter=4),
    }


def pdf_table(headers, rows, widths, styles):
    data = [[Paragraph(html.escape(str(x)), styles["center"]) for x in headers]]
    for row in rows:
        data.append([Paragraph(html.escape(str(x)), styles["ref"]) for x in row])
    table = Table(data, colWidths=[w * cm for w in widths], repeatRows=1, hAlign="CENTER")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DDE8FA")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#081847")),
        ("FONTNAME", (0, 0), (-1, -1), "YaHei"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#AEBBD2")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def draw_pdf_page(canvas, doc):
    canvas.saveState()
    canvas.setFont("YaHei", 7.5)
    canvas.setFillColor(colors.HexColor("#7C8391"))
    canvas.drawRightString(A4[0] - 2.3 * cm, A4[1] - 1.25 * cm, SHORT_NAME + " 项目文档")
    canvas.drawCentredString(A4[0] / 2, 1.15 * cm, str(doc.page))
    canvas.restoreState()


def create_pdf(elements):
    styles = pdf_styles()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, rightMargin=2.1 * cm, leftMargin=2.3 * cm, topMargin=2.0 * cm, bottomMargin=1.8 * cm, title=PROJECT_NAME, author=TEAM_NAME)
    story = [Spacer(1, 2.1 * cm), Paragraph("第八届中国研究生人工智能创新大赛", ParagraphStyle("Cover1", fontName="YaHei", fontSize=19, leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#081847"))), Spacer(1, 2.2 * cm), Paragraph(html.escape(PROJECT_NAME), ParagraphStyle("Cover2", fontName="YaHei", fontSize=24, leading=34, alignment=TA_CENTER, textColor=colors.HexColor("#081847"))), Spacer(1, 0.8 * cm), Paragraph("项目文档", ParagraphStyle("Cover3", fontName="YaHei", fontSize=20, leading=28, alignment=TA_CENTER)), Spacer(1, 2.0 * cm)]
    for text in (VERSION, DATE, TEAM_NAME, GROUP_NAME):
        story.extend([Paragraph(html.escape(text), styles["center"]), Spacer(1, 0.25 * cm)])
    story.extend([PageBreak(), Paragraph("目录", styles["h1"])])
    for line in ["1 项目概况", "1.1 背景和基础", "1.2 场景和价值", "1.3 所需支持", "2 项目规划", "2.1 整体目标", "2.2 技术创新点", "3 实施方案", "3.1 技术可行性分析", "3.2 技术细节", "3.3 计划和分工", "4 参考资料"]:
        story.append(Paragraph(html.escape(line), styles["normal"]))
    story.extend([PageBreak(), Paragraph("记录更改历史", styles["h1"]), pdf_table(["序号", "更改原因", "版本", "作者", "更改日期", "备注"], [["1", "初始编制", VERSION, TEAM_NAME, DATE, "按大赛模板生成"]], [1.1, 3.3, 1.3, 2.8, 2.2, 3.0], styles), PageBreak()])
    for item in elements:
        kind = item[0]
        if kind in {"h1", "h2", "h3"}:
            story.append(Paragraph(html.escape(item[1]), styles[kind]))
        elif kind == "p":
            story.append(Paragraph(html.escape(item[1]), styles["normal"]))
        elif kind == "bullet":
            story.append(Paragraph("• " + html.escape(item[1]), styles["bullet"]))
        elif kind == "formula":
            story.append(Paragraph(html.escape(item[1]), styles["formula"]))
        elif kind == "ref":
            story.append(Paragraph(html.escape(item[1]), styles["ref"]))
        elif kind == "table":
            story.extend([pdf_table(item[1], item[2], item[3], styles), Spacer(1, 0.25 * cm)])
        elif kind == "image":
            img = Image(item[1], width=16.0 * cm, height=8.0 * cm if "architecture" in item[1] else 6.7 * cm)
            story.extend([img, Paragraph(html.escape(item[2]), styles["center"]), Spacer(1, 0.2 * cm)])
        elif kind == "pagebreak":
            story.append(PageBreak())
    doc.build(story, onFirstPage=draw_pdf_page, onLaterPages=draw_pdf_page)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_diagrams()
    elements = build_elements()
    create_docx(elements)
    if PDF_PATH.exists():
        PDF_PATH.unlink()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
